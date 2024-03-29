import os
import re
from datetime import datetime
from difflib import Differ
from docx import Document
from docx.shared import Pt
from Logger import get_logger


class DocxFile:
    full_file: str
    file_name: str
    path: str
    doc_content: Document
    older_content: Document
    bolded_content: Document
    output_doc: Document
    highlighted_content: str

    def __init__(self, path, file):
        self.logger = get_logger(str(self.__class__))
        self.file_name = file
        self.path = path
        self.full_file = os.path.join(self.path, self.file_name + ".docx")
        self.doc_content = None

    def parse(self, filename):
        # Extract docx content
        return Document(os.path.join(self.path, filename))

    def create_new_reference(self):
        if not self.doc_content:
            self.doc_content = self.parse(self.full_file)

        self.write(self.doc_content)

    def highlight_string_content(self, string: str, indexes):
    # Return string with the escape sequences at specific indexes to highlight
        highlight_start = "|"  # TODO: parameterize
        highlight_end = "|"
        words_highlighted = []
        for string_idx, word in enumerate(re.split("([\n ])", string)):
            if string_idx in indexes:
                words_highlighted.append(highlight_start + word + highlight_end)
            else:
                words_highlighted.append(word)
        return "".join(words_highlighted)

    def get_additions_indexes(self, s1: str, s2: str):
    # Return indexes of the additions to s2 compared to s1
        diffs = list(Differ().compare(re.split("([\n ])", s1), re.split("([\n ])", s2)))
        indexes = []
        adj_idx = 0  # Adjust index to compensate for removed words
        for diff_idx, diff in enumerate(diffs):
            if diff[:1] == "+":
                indexes.append(diff_idx - adj_idx)
            elif diff[:1] == "-":
                adj_idx += 1
            elif diff[:1] == "?":
                self.logger.warning("'?' detected! Check the result!")
        return indexes

    def get_paragraphs_as_list(self, content: Document):
        content_listed = []
        for par in content.paragraphs:
            # Exclude title-like portions
            if par.style.name not in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
                # Create a list with all the document's content
                paragraph_string = ""
                for run in par.runs:
                    paragraph_string +=run.text
                content_listed.append(paragraph_string)

        return content_listed

    def compare(self):
        doc_content_listed = self.get_paragraphs_as_list(self.doc_content)
        older_doc_content_listed = self.get_paragraphs_as_list(self.older_content)

        # Create single strings with the documents' content, with newlines as '\n'
        doc_content_string = "\n".join(doc_content_listed)
        older_doc_content_string = "\n".join(older_doc_content_listed)

        compare_indexes = self.get_additions_indexes(older_doc_content_string, doc_content_string)
        self.highlighted_content = self.highlight_string_content(doc_content_string, compare_indexes)

        self.logger.info("Comparison done")

    def core(self):  # TODO: change name
        self.output_doc = Document()

        # First, just copy over some title-like portions, if present
        for doc_par in self.doc_content.paragraphs:
            # TODO: this assumes titles and stuff are only at the beginning of the document
            # so they can just be copied over before everything else
            if doc_par.style.name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
                new_doc_par = self.output_doc.add_paragraph()
                for run in doc_par.runs:
                    new_doc_run = new_doc_par.add_run(run.text)
                    # Run's style data
                    new_doc_run.style = run.style
                    # Run's font data
                    new_doc_run.font.name = "Calibri"
                    new_doc_run.font.size = run.font.size
                # Paragraph's alignment data
                new_doc_par.style = doc_par.style
                new_doc_par.paragraph_format.alignment = doc_par.paragraph_format.alignment

        # Add the rest of the text, with boled additions
        for line in self.highlighted_content.splitlines():
            bolded = re.findall(r"\|[^\|]+\|", line)

            new_doc_par = self.output_doc.add_paragraph()
            if not bolded:
                new_doc_run = new_doc_par.add_run(line)
                # Run's style data
                # new_doc_run.style = run.style # TODO: something useful here?
                # Run's font data
                new_doc_run.font.name = "Courier New"
                new_doc_run.font.size = Pt(11)

            else:  # Some additions present in the line
                # Considering that for each highlighted portion we have two symbols,
                # split() always returns a list with an odd number of elements
                # e.g. "|Simply| trying to |quickly| do something |bold|"
                # -> ['', 'Simply', ' trying to ', 'quickly', ' do something ', 'bold', '']
                portions = line.split("|")
                for index, portion in enumerate(portions):
                    new_doc_run = new_doc_par.add_run(portion)
                    # Make it bold if it was highlighted
                    if index % 2 == 1:
                        new_doc_run.bold = True
                    # Run's font data
                    new_doc_run.font.name = "Courier New"
                    new_doc_run.font.size = Pt(11)

    def write(self, doc: Document):
        filename = self.file_name.split(".")[0] + " - " + datetime.today().strftime("%d-%m-%y") + ".docx"
        new_full_name = os.path.join(self.path, filename)
        self.logger.info('Writing to file "%s"...', {new_full_name})

        doc.save(new_full_name)

        self.logger.info("Done!")

    def run(self, older):
        self.doc_content = self.parse(self.full_file)
        self.older_content = self.parse(older)
        self.compare()
        self.core()
        self.write(self.output_doc)
